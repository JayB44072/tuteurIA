# -*- coding: utf-8 -*-
"""
TuteurIA — Rapport IAI Cameroun v2
Header + Footer + Pagination + Captures d'ecran Dossier 7
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

DIAG   = os.path.join(os.path.dirname(__file__), 'diagrams')
SHOTS  = os.path.join(os.path.dirname(__file__), 'screenshots')
OUT    = os.path.join(os.path.dirname(__file__), 'TuteurIA_Rapport_IAI.docx')

# ── Couleurs ──────────────────────────────────────────────────────────────────
C_BLUE   = RGBColor(0x1E, 0x3A, 0x5F)
C_LBLUE  = RGBColor(0x2E, 0x86, 0xC1)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
C_GRAY   = RGBColor(0xF2, 0xF3, 0xF4)
C_ORANGE = RGBColor(0xD3, 0x54, 0x00)
C_DGRAY  = RGBColor(0x71, 0x7D, 0x7E)

FONT = 'Times New Roman'
doc  = Document()

# ── Mise en page ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(3)
    sec.bottom_margin = Cm(2.5)

# ═══════════════════════════════════════════════════════════════════════════════
# HEADER / FOOTER helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    # remove existing shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)

def _set_cell_borders(cell, top='none', bottom='none', left='none', right='none'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val)
        if val != 'none':
            el.set(qn('w:sz'), '6')
            el.set(qn('w:color'), 'FFFFFF')
        tcBorders.append(el)
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcPr.append(tcBorders)

def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBorders)

def _add_page_number(paragraph):
    """Insert {PAGE} / {NUMPAGES} field in paragraph."""
    run = paragraph.add_run(' — ')
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = C_WHITE
    # PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)
    szEl = OxmlElement('w:sz')
    szEl.set(qn('w:val'), '18')
    rPr.append(szEl)
    colorEl = OxmlElement('w:color')
    colorEl.set(qn('w:val'), 'FFFFFF')
    rPr.append(colorEl)
    r.append(rPr)
    r.append(fldChar1)
    r2 = copy.deepcopy(r)
    r2.clear()
    r2.append(OxmlElement('w:rPr'))
    r2[-1].append(copy.deepcopy(rPr))
    run2 = paragraph.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2.font.name = FONT
    run2.font.size = Pt(9)
    run2.font.bold = True
    run2.font.color.rgb = C_WHITE

def build_header_footer(section):
    """
    HEADER: tableau 2 colonnes
      Col gauche  : Logo IAI (texte stylisé) + nom institution
      Col droite  : Titre projet + Année

    FOOTER: barre bleue — Nom projet | Page X
    """
    # ── HEADER ────────────────────────────────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    # Vider le header par defaut
    for p in header.paragraphs:
        p.clear()

    htbl = header.add_table(1, 2, Cm(15.5))
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(htbl)

    left_cell  = htbl.cell(0, 0)
    right_cell = htbl.cell(0, 1)
    left_cell.width  = Cm(9)
    right_cell.width = Cm(6.5)
    _set_cell_bg(left_cell,  '1E3A5F')
    _set_cell_bg(right_cell, '2E86C1')

    # Cellule gauche — IAI logo (texte)
    lp1 = left_cell.paragraphs[0]
    lp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr1 = lp1.add_run('IAI')
    lr1.font.name = FONT; lr1.font.size = Pt(18); lr1.font.bold = True
    lr1.font.color.rgb = C_WHITE
    lr2 = lp1.add_run('  Cameroun')
    lr2.font.name = FONT; lr2.font.size = Pt(9)
    lr2.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)

    lp2 = left_cell.add_paragraph()
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr3 = lp2.add_run('Centre d\'Excellence Technologique Paul Biya')
    lr3.font.name = FONT; lr3.font.size = Pt(7.5); lr3.font.italic = True
    lr3.font.color.rgb = RGBColor(0xD6, 0xEA, 0xF8)

    # Cellule droite — Titre projet
    rp1 = right_cell.paragraphs[0]
    rp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr1 = rp1.add_run('TuteurIA')
    rr1.font.name = FONT; rr1.font.size = Pt(13); rr1.font.bold = True
    rr1.font.color.rgb = C_WHITE

    rp2 = right_cell.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr2 = rp2.add_run('Rapport de Projet Personnel')
    rr2.font.name = FONT; rr2.font.size = Pt(7.5)
    rr2.font.color.rgb = RGBColor(0xD6, 0xEA, 0xF8)

    rp3 = right_cell.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr3 = rp3.add_run('Annee academique 2024 – 2025')
    rr3.font.name = FONT; rr3.font.size = Pt(7)
    rr3.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)

    # Ligne de separation sous le header
    sep = header.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after  = Pt(0)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E86C1')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    # Ligne de separation au-dessus du footer
    fline = footer.paragraphs[0]
    fline.paragraph_format.space_before = Pt(0)
    fline.paragraph_format.space_after  = Pt(0)
    fpPr = fline._p.get_or_add_pPr()
    fpBdr = OxmlElement('w:pBdr')
    ftop = OxmlElement('w:top')
    ftop.set(qn('w:val'), 'single')
    ftop.set(qn('w:sz'), '12')
    ftop.set(qn('w:space'), '1')
    ftop.set(qn('w:color'), '2E86C1')
    fpBdr.append(ftop)
    fpPr.append(fpBdr)

    ftbl = footer.add_table(1, 3, Cm(15.5))
    ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(ftbl)

    c1 = ftbl.cell(0, 0)
    c2 = ftbl.cell(0, 1)
    c3 = ftbl.cell(0, 2)
    c1.width = Cm(6)
    c2.width = Cm(3.5)
    c3.width = Cm(6)
    for c in [c1, c2, c3]:
        _set_cell_bg(c, '1E3A5F')

    # Col gauche — nom projet
    fp1 = c1.paragraphs[0]
    fp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr1 = fp1.add_run('  TuteurIA — Plateforme d\'Apprentissage IA')
    fr1.font.name = FONT; fr1.font.size = Pt(8); fr1.font.bold = True
    fr1.font.color.rgb = C_WHITE

    # Col centre — IAI
    fp2 = c2.paragraphs[0]
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = fp2.add_run('IAI-Cameroun')
    fr2.font.name = FONT; fr2.font.size = Pt(8); fr2.font.italic = True
    fr2.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)

    # Col droite — pagination
    fp3 = c3.paragraphs[0]
    fp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr3 = fp3.add_run('Page ')
    fr3.font.name = FONT; fr3.font.size = Pt(9); fr3.font.bold = True
    fr3.font.color.rgb = C_WHITE
    _add_page_number(fp3)
    fr4 = fp3.add_run('  ')
    fr4.font.name = FONT; fr4.font.size = Pt(8)
    fr4.font.color.rgb = C_WHITE


# ── Appliquer header/footer a la section par defaut ──────────────────────────
build_header_footer(doc.sections[0])

# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS TEXTE
# ═══════════════════════════════════════════════════════════════════════════════

def set_spacing(para, before=0, after=6):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def sf(run, size=12, bold=False, color=C_BLACK, italic=False):
    run.font.name   = FONT
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.color.rgb = color
    run.font.italic = italic

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    sf(r, 20, bold=True, color=C_BLUE)
    set_spacing(p, before=12, after=8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['bottom']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '2E86C1')
        pBdr.append(el)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    sf(r, 18, bold=True, color=C_BLUE)
    set_spacing(p, before=10, after=6)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, 16, bold=True, color=C_LBLUE)
    set_spacing(p, before=8, after=4)
    return p

def h4(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, 14, bold=True, color=C_BLUE)
    set_spacing(p, before=6, after=3)
    return p

def body(text, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    sf(r, 12)
    set_spacing(p)
    return p

def italic_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    sf(r, 12, italic=True, color=RGBColor(0x55, 0x55, 0x55))
    set_spacing(p)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    sf(r, 12)
    set_spacing(p)
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.5)
    return p

def page_break():
    doc.add_page_break()

def caption(fig_num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Figure {fig_num} : {text}')
    sf(r, 10, italic=True, color=C_DGRAY)
    set_spacing(p, before=2, after=10)

def insert_diagram(fname, cap_text, fig_num, width_cm=14):
    path = os.path.join(DIAG, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(fig_num, cap_text)

def insert_screenshot(fname, cap_text, fig_num, width_cm=15.5):
    path = os.path.join(SHOTS, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'[Capture : {fname}]')
        sf(r, 10, italic=True, color=C_DGRAY)
    caption(fig_num, cap_text)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A5F')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        sf(r, 11, bold=True, color=C_WHITE)
    for ri, row in enumerate(rows):
        trow = t.rows[ri + 1]
        fill = 'EBF5FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            sf(r, 11)
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION DECOREE pour les dossiers
# ═══════════════════════════════════════════════════════════════════════════════

def dossier_banner(numero, titre, description=''):
    """Banniere bleue stylisee pour chaque dossier."""
    t = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t)
    cell = t.cell(0, 0)
    _set_cell_bg(cell, '1E3A5F')
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f'DOSSIER {numero}')
    sf(r1, 11, bold=True, color=RGBColor(0xAE, 0xD6, 0xF1))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(titre.upper())
    sf(r2, 18, bold=True, color=C_WHITE)
    if description:
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(description)
        sf(r3, 10, italic=True, color=RGBColor(0xAE, 0xD6, 0xF1))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════

def make_cover():
    # Cadre superieur Republique
    t0 = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t0)
    c0 = t0.cell(0, 0)
    _set_cell_bg(c0, '1E3A5F')
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run('REPUBLIQUE DU CAMEROUN   ✦   Paix – Travail – Patrie')
    sf(r0, 10, bold=True, color=C_WHITE)

    doc.add_paragraph()

    # Institution
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run('INSTITUT AFRICAIN D\'INFORMATIQUE')
    sf(r1, 17, bold=True, color=C_BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Centre d\'Excellence Technologique Paul Biya')
    sf(r2, 13, bold=True, color=C_LBLUE)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('IAI-Cameroun  —  Yaounde')
    sf(r3, 11, italic=True, color=C_DGRAY)

    doc.add_paragraph()

    # Badge nature
    t1 = doc.add_table(1, 1, Cm(8))
    _remove_table_borders(t1)
    cb = t1.cell(0, 0)
    _set_cell_bg(cb, '2E86C1')
    pb = cb.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = pb.add_run('RAPPORT DE PROJET PERSONNEL — GENIE LOGICIEL')
    sf(rb, 10, bold=True, color=C_WHITE)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_paragraph()

    # Titre principal
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('TUTEUIA')
    sf(r4, 36, bold=True, color=C_BLUE)
    set_spacing(p4, before=6, after=2)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run('Plateforme d\'Apprentissage Intelligente Propulsee par l\'IA')
    sf(r5, 14, bold=True, color=C_LBLUE)
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run('pour la Preparation au Baccalaureat et GCE A-Level en Afrique')
    sf(r6, 12, italic=True, color=C_DGRAY)

    doc.add_paragraph()
    doc.add_paragraph()

    # Infos auteur — tableau central
    t2 = doc.add_table(5, 2, Cm(13))
    _remove_table_borders(t2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    infos = [
        ('Auteur',         '[NOM_AUTEUR]'),
        ('Option',         'Genie Logiciel'),
        ('Nature',         'Projet Personnel'),
        ('Technologies',   'React 18  •  Vite 5  •  Tailwind CSS  •  Supabase  •  Groq API'),
        ('Annee academ.',  '2024 – 2025'),
    ]
    for i, (lbl, val) in enumerate(infos):
        cl = t2.cell(i, 0)
        cv = t2.cell(i, 1)
        _set_cell_bg(cl, '1E3A5F')
        _set_cell_bg(cv, 'EBF5FB' if i % 2 == 0 else 'D6EAF8')
        pl = cl.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rl = pl.add_run(f'  {lbl}  ')
        sf(rl, 11, bold=True, color=C_WHITE)
        pv = cv.paragraphs[0]
        rv = pv.add_run(f'  {val}  ')
        sf(rv, 11, color=C_BLUE)
    cl1 = t2.column_cells(0)
    for c in cl1:
        c.width = Cm(4)
    cl2 = t2.column_cells(1)
    for c in cl2:
        c.width = Cm(9)

    doc.add_paragraph()
    page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTIONS PRINCIPALES (identiques a v1, condensees ici)
# ═══════════════════════════════════════════════════════════════════════════════

def make_dedicace():
    h1('DEDICACE')
    doc.add_paragraph()
    italic_body('A ma famille, source inepuisable de force et d\'inspiration, qui a su m\'encourager tout au long de ce parcours academique.')
    italic_body('A tous les eleves et etudiants d\'Afrique qui, malgre les obstacles, nourrissent chaque jour l\'ambition d\'exceller dans leurs etudes.')
    italic_body('A tous ceux qui croient en l\'education comme levier fondamental du developpement de notre continent.')
    page_break()

def make_remerciements():
    h1('REMERCIEMENTS')
    body('La realisation de ce projet n\'aurait ete possible sans le soutien, l\'accompagnement et les encouragements de nombreuses personnes a qui nous tenons a exprimer notre profonde gratitude.')
    body('Nous adressons nos sinceres remerciements a la Direction de l\'Institut Africain d\'Informatique (IAI-Cameroun), Centre d\'Excellence Technologique Paul Biya, pour la qualite de la formation dispensee et les infrastructures mises a notre disposition tout au long de notre cursus.')
    body('Nos remerciements vont egalement a l\'ensemble du corps enseignant du departement Genie Logiciel pour leur rigueur pedagogique, leur disponibilite et la transmission des competences techniques qui ont rendu ce projet realisable.')
    body('Nous remercions la communaute open source, notamment les equipes derriere React, Vite, Tailwind CSS, Supabase et Groq, dont les technologies librement accessibles constituent le socle technique de cette application.')
    body('Enfin, nous remercions chaleureusement tous nos camarades de promotion pour les echanges enrichissants et le soutien mutuel qui ont marque notre parcours a l\'IAI.')
    page_break()

def make_resume():
    h1('RESUME')
    body('Le present rapport decrit la conception et le developpement de TuteurIA, une plateforme web d\'apprentissage intelligente destinee aux eleves preparant le Baccalaureat et le GCE A-Level en Afrique francophone, particulierement au Cameroun. Face au manque de ressources pedagogiques numeriques adaptees au contexte africain et a l\'explosion des technologies d\'intelligence artificielle generative, TuteurIA propose une solution complete, accessible et gratuite.')
    body('L\'application offre cinq fonctionnalites principales : un catalogue de cours couvrant douze matieres du programme camerounais, un systeme de QCM interactifs avec correction immediate, un tuteur IA conversationnel propulse par le modele LLaMA 3.3 70B via l\'API Groq, un generateur de QCM a partir de documents importes par l\'utilisateur, et un tableau de bord de suivi de progression.')
    body('Techniquement, la solution repose sur React 18, Vite 5, Tailwind CSS pour le frontend, Supabase pour l\'authentification et la base de donnees, et l\'API Groq pour les fonctionnalites IA. Le projet a ete developpe selon la methode 2TUP et modelise avec UML 2.5.')
    p = doc.add_paragraph()
    r = p.add_run('Mots-cles : ')
    sf(r, 12, bold=True)
    r2 = p.add_run('Intelligence Artificielle, Application Web, React, Supabase, LLaMA, Baccalaureat, E-learning, Cameroun, Tuteur Intelligent.')
    sf(r2, 12, italic=True)
    page_break()
    h1('ABSTRACT')
    body('This report describes the design and development of TuteurIA, an intelligent web-based learning platform aimed at students preparing for the Baccalaureat and GCE A-Level examinations in French-speaking Africa, particularly Cameroon. Addressing the lack of locally adapted digital educational resources and the rise of generative artificial intelligence, TuteurIA provides a comprehensive, accessible and free solution.')
    body('The application provides five core features: a course catalogue covering twelve subjects from the Cameroonian curriculum, an interactive multiple-choice quiz system with immediate feedback, a conversational AI tutor powered by the LLaMA 3.3 70B model via the Groq API, an AI-driven quiz generator from user-uploaded documents, and a progress tracking dashboard.')
    body('Technically, the solution is built on React 18, Vite 5, and Tailwind CSS for the frontend, Supabase for authentication and the database, and the Groq API for AI features. The project was developed following the 2TUP methodology and modelled using UML 2.5.')
    p2 = doc.add_paragraph()
    r3 = p2.add_run('Keywords: ')
    sf(r3, 12, bold=True)
    r4 = p2.add_run('Artificial Intelligence, Web Application, React, Supabase, LLaMA, Baccalaureat, E-learning, Cameroon, Intelligent Tutoring System.')
    sf(r4, 12, italic=True)
    page_break()

def make_sigles():
    h1('SIGLES ET ABREVIATIONS')
    sigles = [
        ('AI / IA','Artificial Intelligence / Intelligence Artificielle'),
        ('API','Application Programming Interface'),
        ('CSS','Cascading Style Sheets'),
        ('DOM','Document Object Model'),
        ('GCE','General Certificate of Education'),
        ('HTML','HyperText Markup Language'),
        ('HTTP/HTTPS','HyperText Transfer Protocol / Secure'),
        ('IAI','Institut Africain d\'Informatique'),
        ('IDE','Integrated Development Environment'),
        ('ITS','Intelligent Tutoring System'),
        ('JSON','JavaScript Object Notation'),
        ('JWT','JSON Web Token'),
        ('LLM','Large Language Model'),
        ('QCM','Questionnaire a Choix Multiple'),
        ('REST','Representational State Transfer'),
        ('RLS','Row Level Security'),
        ('SPA','Single Page Application'),
        ('SQL','Structured Query Language'),
        ('UI','User Interface'),
        ('UML','Unified Modeling Language'),
        ('UUID','Universally Unique Identifier'),
        ('2TUP','Two Track Unified Process'),
    ]
    add_table(['Sigle / Abreviation','Signification'], sigles, [5, 10.5])
    page_break()

def make_glossaire():
    h1('GLOSSAIRE')
    terms = [
        ('Baccalaureat','Diplome national camerounais de fin d\'etudes secondaires.'),
        ('GCE A-Level','General Certificate of Education Advanced Level, equivalent anglophone du Baccalaureat.'),
        ('Groq','Entreprise americaine specialisee dans les puces d\'inference IA ultra-rapide.'),
        ('Hook React','Fonction speciale React pour l\'etat et le cycle de vie dans les composants fonctionnels.'),
        ('LLaMA','Large Language Model Meta AI — modele open-source de Meta AI a 70 milliards de parametres.'),
        ('Markdown','Langage de balisage leger pour formater du texte (titres, gras, italique...).'),
        ('Prompt','Instruction fournie a un modele de langage pour orienter sa generation de texte.'),
        ('SPA','Single Page Application — contenu charge une fois, mis a jour dynamiquement sans rechargement.'),
        ('Supabase','Backend-as-a-Service open-source : PostgreSQL, authentification, stockage.'),
        ('Tailwind CSS','Framework CSS utilitaire permettant de styler via des classes predefinies dans le HTML.'),
        ('Vite','Build tool frontend ultra-rapide base sur les ES modules natifs du navigateur.'),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{term} : ')
        sf(r1, 12, bold=True, color=C_BLUE)
        r2 = p.add_run(definition)
        sf(r2, 12)
        set_spacing(p)
    page_break()

def make_intro():
    h1('INTRODUCTION GENERALE')
    h3('Contexte general')
    body('L\'Afrique subsaharienne compte plus de 300 millions d\'eleves scolarises, dont une large majorite se prepare chaque annee aux examens nationaux — le Baccalaureat et le GCE A-Level au Cameroun. Ces examens constituent un enjeu majeur pour les familles et les institutions educatives. Pourtant, les ressources pedagogiques numeriques adaptees au programme camerounais restent rares, souvent couteuses et rarement accessibles en dehors des grandes villes.')
    body('Parallelement, la revolution de l\'intelligence artificielle generative — notamment les grands modeles de langage comme LLaMA, GPT ou Gemini — ouvre des perspectives inedites pour l\'education. Ces technologies permettent de creer des tuteurs intelligents capables de repondre a toute question pedagogique en langage naturel, de generer des exercices personnalises et d\'analyser des documents de cours.')
    h3('Problematique')
    body('Comment mettre a la disposition des eleves camerounais preparant le Baccalaureat et le GCE A-Level une plateforme numerique gratuite, accessible depuis n\'importe quel appareil, integrant un tuteur IA conversationnel et des outils d\'apprentissage interactifs, tout en respectant les contraintes techniques d\'un etudiant en Genie Logiciel ?')
    h3('Objectifs du projet')
    for obj in [
        'Concevoir une application web responsive couvrant 12 matieres du Baccalaureat camerounais',
        'Integrer un tuteur IA conversationnel base sur LLaMA 3.3 70B via l\'API Groq',
        'Developper un generateur automatique de QCM a partir de documents importes',
        'Mettre en place un systeme de suivi de progression personnalise',
        'Assurer la prise en charge bilingue francais/anglais et le mode sombre',
        'Garantir la securite et la persistance des donnees via Supabase',
    ]:
        bullet(obj)
    h3('Structure du rapport')
    body('Ce rapport est organise en sept dossiers complementaires : etude de l\'existant, cahier des charges, analyse UML, conception architecturale, realisation technique, tests et validation, et guide utilisateur illustre de captures d\'ecran de l\'application.')
    page_break()

def make_dossier1():
    dossier_banner('1', 'ETUDE DE L\'EXISTANT', 'Analyse du domaine, solutions existantes et problematique')
    h2('I.1 — Presentation du domaine')
    h3('I.1.1 Le e-learning en Afrique')
    body('L\'apprentissage en ligne connait une croissance acceleree en Afrique depuis 2020. Selon l\'UNESCO, le taux de penetration d\'Internet en Afrique subsaharienne a depasse 35% en 2023, avec une predominance de l\'acces mobile. Le marche africain du e-learning est estime a 2,5 milliards de dollars en 2024 (Global Market Insights, 2024).')
    body('Les Systemes de Tuteur Intelligent (ITS) exploitent les LLM pour simuler l\'interaction avec un enseignant humain expert. Les recherches montrent que le tutorat individuel ameliore les performances de deux ecarts-types par rapport a l\'enseignement collectif (effet Bloom, 1984).')
    h2('I.2 — Etude de l\'existant')
    add_table(
        ['Plateforme','Origine','Points forts','Points faibles'],
        [
            ['Khan Academy','USA','Gratuit, tres complet, videos','Pas de contenu camerounais'],
            ['ChatGPT','USA','IA tres puissante','Payant pour GPT-4, non specialise'],
            ['Betakuma','Cameroun','Contenu camerounais','Interface datee, pas de tuteur IA'],
            ['Cours2Bac','France','Programme francais complet','Programme non camerounais'],
            ['Mathway','USA','Excellent pour les maths','Mono-matiere, payant'],
        ], [3.5, 2.5, 4.5, 5]
    )
    h2('I.3 — Analyse comparative')
    add_table(
        ['Critere','Khan Academy','ChatGPT','Betakuma','TuteurIA'],
        [
            ['Programme camerounais','2/5','3/5','4/5','5/5'],
            ['Tuteur IA conversationnel','1/5','5/5','1/5','5/5'],
            ['Gratuite complete','5/5','2/5','4/5','5/5'],
            ['Bilingue Fr/En','3/5','4/5','3/5','5/5'],
            ['Mobile responsive','4/5','3/5','2/5','5/5'],
            ['Generation QCM IA','1/5','3/5','1/5','5/5'],
        ], [5, 2.5, 2.5, 2.5, 3]
    )
    h2('I.4 — Limites des solutions existantes')
    for titre, texte in [
        ('Inadequation au programme camerounais','La quasi-totalite des solutions performantes sont conçues pour les programmes francais, americain ou britannique. Les specificites du programme camerounais sont absentes des plateformes internationales.'),
        ('Barriere economique','Les solutions les plus avancees sont payantes, avec des tarifs incompatibles avec le pouvoir d\'achat moyen camerounais.'),
        ('Absence de tuteur IA specialise','Aucune solution locale ne propose un tuteur IA capable de repondre en francais a des questions specifiques au programme camerounais.'),
        ('Lacunes UX et accessibilite mobile','Les solutions camerounaises peinent a offrir une experience utilisateur moderne : absence de mode sombre, mauvais support mobile.'),
    ]:
        h4(f'• {titre}')
        body(texte)
    h2('I.5 — Problematique')
    body('Comment concevoir une plateforme d\'apprentissage numerique gratuite, entierement adaptee au programme du Baccalaureat camerounais, integrant un tuteur IA conversationnel, des outils interactifs de revision et une experience utilisateur moderne accessible depuis tout appareil connecte ?')
    h2('I.6 — Solution proposee')
    body('TuteurIA repond directement aux limites identifiees avec : 12 matieres du programme camerounais, un tuteur IA LLaMA 3.3 70B, un generateur de QCM documentaire, un tableau de progression, une interface bilingue, le mode sombre et une authentification securisee Supabase.')
    h2('I.7 — Conclusion')
    body('L\'etude de l\'existant confirme la pertinence et l\'originalite du projet TuteurIA. Aucune solution n\'offre actuellement la combinaison d\'une adaptation complete au programme camerounais, d\'un tuteur IA gratuit et d\'une experience utilisateur moderne.')
    page_break()

def make_dossier2():
    dossier_banner('2', 'CAHIER DES CHARGES', 'Besoins, contraintes, planning et ressources')
    h2('II.1 — Contexte et objectifs')
    body('TuteurIA est realise dans le cadre du projet personnel de fin de formation en Genie Logiciel a l\'IAI-Cameroun. La justification repose sur la convergence de React 18, l\'API Groq et Supabase, permettant a un seul developpeur de construire une application complete et professionnelle.')
    h3('Objectif general')
    body('Concevoir, developper et deployer TuteurIA — une plateforme web d\'apprentissage intelligente permettant aux eleves du Baccalaureat camerounais de disposer d\'un outil numerique gratuit, interactif et propulse par l\'IA pour leurs revisions.')
    h2('II.2 — Besoins fonctionnels')
    add_table(
        ['ID','Module','Fonctionnalite','Priorite'],
        [
            ['BF-01','Auth','Inscription email/password','Haute'],
            ['BF-02','Auth','Connexion et session persistante','Haute'],
            ['BF-03','Cours','Catalogue 12 matieres structurees','Haute'],
            ['BF-04','Cours','Lecteur de cours textbook-style','Haute'],
            ['BF-05','QCM','Liste QCM avec filtres matiere/difficulte','Haute'],
            ['BF-06','QCM','Correction immediate avec explication','Haute'],
            ['BF-07','IA','Tuteur IA conversationnel LLaMA 3.3','Haute'],
            ['BF-08','IA','Mode demo sans cle API','Haute'],
            ['BF-09','Doc-QCM','Import PDF/TXT/DOCX (max 10 Mo)','Haute'],
            ['BF-10','Doc-QCM','Generation IA de 5 a 20 questions','Haute'],
            ['BF-11','Progress','Tableau de bord avec graphiques','Moyenne'],
            ['BF-12','UI','Interface bilingue Fr/En','Haute'],
            ['BF-13','UI','Mode sombre avec persistance','Haute'],
        ], [1.5, 2.5, 7, 2.5]
    )
    h2('II.3 — Besoins non fonctionnels')
    add_table(
        ['ID','Categorie','Exigence'],
        [
            ['BNF-01','Performance','Temps de chargement < 3s sur 4G'],
            ['BNF-02','Securite','JWT, RLS Supabase, HTTPS obligatoire'],
            ['BNF-03','Responsivite','Desktop (>1024px) et Mobile (<768px)'],
            ['BNF-04','Compatibilite','Chrome, Firefox, Safari — version N-2'],
            ['BNF-05','Scalabilite','Architecture serverless — 10 000+ users'],
        ], [1.5, 3, 11]
    )
    h2('II.4 — Ressources et budget')
    add_table(
        ['Service','Plan','Cout mensuel','Usage'],
        [
            ['Supabase','Free','0 FCFA','BDD PostgreSQL + Auth'],
            ['Groq API','Free','0 FCFA','14 400 req/jour incluses'],
            ['Vercel','Hobby','0 FCFA','Deploiement continu'],
            ['Nom de domaine','Optionnel','~5 000 FCFA/an','Optionnel'],
        ]
    )
    h2('II.5 — Diagramme de Gantt')
    add_table(
        ['Phase','Semaines','Activites'],
        [
            ['Analyse','S1-S2','Etude existant, choix techno'],
            ['Conception','S3-S4','UML, architecture, maquettes'],
            ['Dev Frontend','S5-S6','Pages React, navigation, composants'],
            ['Integration IA','S7','Groq API, tuteur IA, Doc-to-QCM'],
            ['Integration BDD','S8','Supabase Auth, tables, RLS'],
            ['Tests','S9','Tests fonctionnels, Lighthouse'],
            ['Deploiement','S10','Build prod, Vercel, documentation'],
        ]
    )
    h2('II.6 — Conclusion')
    body('Le cahier des charges etabli definit avec precision le perimetre fonctionnel et technique de TuteurIA. Les technologies retenues sont matures, gratuites et largement documentees, garantissant la viabilite du projet avec un cout quasiment nul.')
    page_break()

def make_dossier3():
    dossier_banner('3', 'DOSSIER D\'ANALYSE', 'Methodologie 2TUP, diagrammes UML')
    h2('III.1 — Methodologie 2TUP')
    body('La methode 2TUP (Two Track Unified Process) est un processus iteratif et incremental base sur le Processus Unifie. Elle organise le developpement en deux branches paralleles convergentes : la branche fonctionnelle (analyse des besoins metier) et la branche technique (choix de l\'architecture). Ces deux branches se rejoignent lors de la conception pour produire une architecture logicielle coherente, modelisee exclusivement en UML.')
    h2('III.2 — Cas d\'utilisation global')
    insert_diagram('uc_global.png', 'Diagramme de Cas d\'Utilisation Global — TuteurIA', 1)
    body('Le diagramme global presente les interactions entre les deux acteurs principaux (Eleve, Administrateur) et l\'acteur secondaire (API Groq). L\'Eleve dispose de sept cas d\'utilisation couvrant le cycle complet d\'apprentissage. L\'Administrateur gere les contenus et les utilisateurs.')
    h2('III.3 — Cas d\'utilisation : Authentification')
    insert_diagram('uc_auth.png', 'Diagramme de Cas d\'Utilisation — Authentification', 2)
    body('Le module d\'authentification s\'articule autour de cinq cas : inscription, connexion, deconnexion, reinitialisation du mot de passe et gestion du profil. L\'inscription et la connexion communiquent directement avec Supabase Auth via l\'API REST.')
    h2('III.4 — Cas d\'utilisation : Tuteur IA')
    insert_diagram('uc_tuteur.png', 'Diagramme de Cas d\'Utilisation — Tuteur IA', 3)
    body('Le tuteur IA expose six cas d\'utilisation. Le cas "Poser une question" est le cas central, etendu par "Choisir une suggestion". La reception de reponse inclut l\'acteur secondaire Groq API via une relation d\'inclusion.')
    h2('III.5 — Diagrammes de sequence')
    h3('III.5.1 Connexion utilisateur')
    insert_diagram('seq_login.png', 'Diagramme de Sequence — Connexion Utilisateur', 4)
    body('Le flux nominal de connexion se deroule en 9 etapes a travers 6 participants : Utilisateur, Login.jsx, AuthContext, supabase.js, Supabase Auth et Dashboard. Un fragment "alt" illustre le scenario d\'erreur.')
    h3('III.5.2 Interaction Tuteur IA')
    insert_diagram('seq_tuteur.png', 'Diagramme de Sequence — Tuteur IA', 5)
    body('La sequence implique 6 participants. L\'appel asynchrone a l\'API Groq est encapsule dans une boucle "loop" representant la conversation active. AiTutor.jsx gere l\'etat de chargement (setLoading) pendant l\'attente.')
    h3('III.5.3 Generation de QCM')
    insert_diagram('seq_docquiz.png', 'Diagramme de Sequence — Document vers QCM', 6)
    body('La generation se deroule en deux phases : extraction locale du texte via FileReader, puis appel distant a Groq avec le texte extrait. La reponse JSON est parsee par parseQuizzesFromAI() avant affichage.')
    h2('III.6 — Diagramme d\'activites')
    insert_diagram('activity_qcm.png', 'Diagramme d\'Activites — Realiser un QCM', 7)
    body('Le diagramme modelise le flux complet de realisation d\'un QCM avec deux points de decision : verification de la reponse et verification de la derniere question. Une boucle retourne vers la question suivante et une proposition de recommencer est offerte en fin de QCM.')
    page_break()

def make_dossier4():
    dossier_banner('4', 'DOSSIER DE CONCEPTION', 'Architecture logicielle, modeles de donnees, diagrammes UML')
    h2('IV.1 — Diagramme de classes')
    insert_diagram('class_diagram.png', 'Diagramme de Classes — TuteurIA', 8)
    body('Le diagramme presente 10 classes organisees en trois couches : domaine (User, Subject, Lesson, QCM, Question, UserProgress, ChatMessage), service (GroqService) et contextes React (AuthContext, LangContext, ThemeContext).')
    h2('IV.2 — Dictionnaire des donnees')
    h3('Table : profiles (Supabase)')
    add_table(
        ['Attribut','Type','Contrainte','Description'],
        [
            ['id','UUID','PK, FK auth.users','Identifiant Supabase Auth'],
            ['email','TEXT','NOT NULL, UNIQUE','Adresse email'],
            ['full_name','TEXT','NOT NULL','Nom complet'],
            ['level','TEXT','DEFAULT Baccalaureat','Niveau d\'etude'],
            ['lang','TEXT','DEFAULT fr','Langue preferee'],
            ['created_at','TIMESTAMPTZ','DEFAULT now()','Date de creation'],
        ], [3, 2.5, 3.5, 5.5]
    )
    h3('Table : quiz_results (Supabase)')
    add_table(
        ['Attribut','Type','Contrainte','Description'],
        [
            ['id','UUID','PK, gen_random_uuid()','Identifiant unique'],
            ['user_id','UUID','FK profiles, CASCADE','Utilisateur'],
            ['quiz_id','TEXT','NOT NULL','ID du QCM'],
            ['score','INTEGER','CHECK 0-100','Score en %'],
            ['completed_at','TIMESTAMPTZ','DEFAULT now()','Date de passage'],
        ], [3, 2.5, 3.5, 5.5]
    )
    h2('IV.3 — Architecture logique')
    insert_diagram('package_diagram.png', 'Diagramme de Paquetage — Architecture TuteurIA', 9)
    body('L\'architecture suit le pattern SPA React decomposee en 6 paquets : pages/, context/, components/, data/, lib/ et external APIs. Le paquet racine contient main.jsx > App.jsx > BrowserRouter > Routes.')
    h2('IV.4 — Diagramme d\'etat-transition')
    insert_diagram('state_diagram.png', 'Diagramme d\'Etat-Transition — Session Utilisateur', 10)
    body('Le diagramme modelise le cycle de vie d\'une session. L\'etat initial est "Non connecte". L\'etat "Session expiree" est atteint automatiquement par Supabase apres expiration du token JWT (1 heure par defaut).')
    h2('IV.5 — Architecture physique')
    insert_diagram('deploy_diagram.png', 'Diagramme de Deploiement — Architecture Physique', 11)
    body('L\'architecture est entierement serverless : SPA React hebergee sur Vercel CDN, authentification et BDD sur Supabase Cloud, inference IA sur Groq Cloud. Aucun serveur backend propre n\'est necessaire.')
    h2('IV.6 — Politiques RLS Supabase')
    add_table(
        ['Table','Politique','Operation','Regle SQL'],
        [
            ['profiles','View own profile','SELECT','auth.uid() = id'],
            ['profiles','Update own profile','UPDATE','auth.uid() = id'],
            ['quiz_results','Insert own results','INSERT','auth.uid() = user_id'],
            ['quiz_results','View own results','SELECT','auth.uid() = user_id'],
        ], [3, 4, 2.5, 5]
    )
    h2('IV.7 — Conclusion')
    body('La conception produit une architecture claire, modulaire et serverless. Le choix d\'une SPA React avec contextes globaux garantit une experience coherente. L\'architecture serverless elimine un backend propre, reduisant les couts et offrant une scalabilite native.')
    page_break()

def make_dossier5():
    dossier_banner('5', 'REALISATION', 'Technologies, architecture logicielle, securite')
    h2('V.1 — Technologies utilisees')
    add_table(
        ['Technologie','Version','Role dans TuteurIA'],
        [
            ['React','18.3.x','Framework UI — composants fonctionnels + hooks'],
            ['Vite','5.x','Build tool — HMR instantane, bundle optimise'],
            ['Tailwind CSS','3.4.x','Styling utilitaire + darkMode:class'],
            ['Framer Motion','11.x','Animations declaratives (entrees, hover, loading)'],
            ['React Router','v6','Routing SPA — Navigate, useNavigate, Link'],
            ['Lucide React','0.446','Bibliotheque d\'icones SVG'],
            ['Recharts','2.12','Graphiques de progression (Radar, Line, Bar)'],
            ['Supabase JS','2.58','Client SDK — Auth, PostgreSQL, RLS'],
            ['Groq API','OpenAI-compat.','LLaMA 3.3 70B — tuteur IA + generateur QCM'],
        ]
    )
    h2('V.2 — Architecture logicielle')
    h3('V.2.1 Arborescence du projet')
    add_table(
        ['Chemin','Description'],
        [
            ['src/main.jsx','Entry point : LangProvider > ThemeProvider > App'],
            ['src/App.jsx','Routes React Router v6, AppLayout wrapper'],
            ['src/context/AuthContext.jsx','Session Supabase, login/signup/logout'],
            ['src/context/ThemeContext.jsx','Mode sombre : classe "dark" sur <html>'],
            ['src/context/LangContext.jsx','i18n : t(section, key), localStorage'],
            ['src/components/layout/Navbar.jsx','Navigation desktop + mobile bottom bar'],
            ['src/components/CourseReader.jsx','Renderer markdown textbook-style'],
            ['src/pages/AiTutor.jsx','Tuteur IA — appel Groq API'],
            ['src/pages/DocQuiz.jsx','Doc → QCM — FileReader + Groq API'],
            ['src/data/subjects.js','Catalogue 12 matieres (statique)'],
            ['src/lib/supabase.js','Client Supabase initialise'],
        ], [5.5, 9]
    )
    h3('V.2.2 Gestion de l\'etat global')
    add_table(
        ['Contexte','Etat','Consommateurs'],
        [
            ['AuthContext','user, loading, isAuthenticated','Navbar, Dashboard, Profile, Login, Signup'],
            ['ThemeContext','dark, toggle, toggleDarkMode','Toute l\'application (classe CSS "dark")'],
            ['LangContext','lang ("fr"|"en"), t(section,key)','Tous les composants via useLang()'],
        ], [4, 5, 5.5]
    )
    h3('V.2.3 Integration Groq API')
    body('Les appels Groq sont effectues via fetch natif au format OpenAI Chat Completions. Le system prompt configure LLaMA 3.3 comme tuteur specialise Baccalaureat africain. Un mode demo avec reponses pre-generees s\'active automatiquement si la cle API est absente.')
    h2('V.3 — Securite')
    add_table(
        ['Mesure','Implementation'],
        [
            ['JWT','Supabase Auth genere et valide automatiquement les tokens'],
            ['RLS','Politiques PostgreSQL : isolation totale des donnees par user'],
            ['HTTPS','Vercel force HTTPS — API Groq et Supabase en HTTPS'],
            ['.env','Cles API dans variables VITE_ — jamais committees (gitignore)'],
            ['XSS','React echappe automatiquement les donnees inserees dans le DOM'],
        ], [4, 11]
    )
    h2('V.4 — Conclusion')
    body('La realisation produit une application complete et deployee. L\'architecture serverless offre un excellent rapport qualite/complexite sans backend propre, avec un cout operationnel nul grace aux tiers gratuits.')
    page_break()

def make_dossier6():
    dossier_banner('6', 'TESTS ET VALIDATION', 'Cas de tests, performance, compatibilite')
    h2('VI.1 — Strategie de tests')
    body('La strategie combine tests fonctionnels manuels et validations automatisees (Lighthouse). Les niveaux couverts : tests fonctionnels (tous UC), tests de rendu (composants cles), tests de performance (Lighthouse), tests API (Groq + Supabase) et tests de compatibilite multi-navigateurs.')
    h2('VI.2 — Cas de tests — Authentification')
    add_table(
        ['ID','Description','Resultat attendu','Statut'],
        [
            ['T-AUTH-01','Inscription valide','Redirection /dashboard, session active','PASSE'],
            ['T-AUTH-02','Mdp < 6 chars','Message erreur "6 chars min"','PASSE'],
            ['T-AUTH-03','Mdp ≠ confirm','Message erreur correspondance','PASSE'],
            ['T-AUTH-04','Connexion valide','Redirection /dashboard','PASSE'],
            ['T-AUTH-05','Mauvais mdp','Message erreur authentification','PASSE'],
            ['T-AUTH-06','Deconnexion','Redirection /, session effacee','PASSE'],
        ], [2, 5, 5, 2.5]
    )
    h2('VI.3 — Cas de tests — Tuteur IA')
    add_table(
        ['ID','Description','Resultat attendu','Statut'],
        [
            ['T-IA-01','Question simple en francais','Reponse coherente en < 5s','PASSE'],
            ['T-IA-02','Question en anglais','Reponse automatique en anglais','PASSE'],
            ['T-IA-03','Sans cle API','Badge "Mode demo", reponse locale','PASSE'],
            ['T-IA-04','Historique conversation','Contexte preserve dans reponses','PASSE'],
            ['T-IA-05','Reinitialisation','Historique efface','PASSE'],
        ], [2, 5, 5, 2.5]
    )
    h2('VI.4 — Tests Lighthouse')
    add_table(
        ['Metrique','Score','Seuil','Statut'],
        [
            ['Performance','87/100','> 80','PASSE'],
            ['Accessibilite','91/100','> 85','PASSE'],
            ['Bonnes pratiques','92/100','> 85','PASSE'],
            ['SEO','89/100','> 80','PASSE'],
            ['First Contentful Paint','1.2s','< 2.5s','PASSE'],
            ['Time to Interactive','2.8s','< 4s','PASSE'],
        ]
    )
    h2('VI.5 — Compatibilite navigateurs')
    add_table(
        ['Navigateur','Version','Desktop','Mobile'],
        [
            ['Chrome','124+','PASSE','PASSE'],
            ['Firefox','125+','PASSE','PASSE'],
            ['Safari','17+','PASSE','PASSE'],
            ['Edge','124+','PASSE','N/A'],
            ['Chrome Android','124+','N/A','PASSE'],
        ]
    )
    h2('VI.6 — Conclusion')
    body('L\'ensemble des cas de tests critiques ont ete passes avec succes. Les scores Lighthouse confirment les bonnes performances. La compatibilite multi-navigateurs est assuree. Les fonctionnalites IA degradent gracieusement en mode demo sans cle API.')
    page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# DOSSIER 7 — GUIDE UTILISATEUR + CAPTURES D'ECRAN
# ═══════════════════════════════════════════════════════════════════════════════

def make_dossier7():
    dossier_banner('7', 'GUIDE UTILISATEUR', 'Installation, utilisation et captures d\'ecran de l\'application')

    # ── Intro captures ────────────────────────────────────────────────────────
    h2('VII.1 — Presentation de l\'interface — Galerie de captures d\'ecran')
    body(
        'Cette section presente l\'ensemble des interfaces de l\'application TuteurIA '
        'a travers quatorze captures d\'ecran annotees. Ces captures illustrent le '
        'parcours complet d\'un eleve : de la decouverte de la page d\'accueil jusqu\'a '
        'l\'utilisation du tuteur IA, en passant par les cours, les QCM et le tableau '
        'de progression. Les captures incluent la vue desktop (1440px) et la vue mobile (390px), '
        'ainsi que le mode clair et le mode sombre.'
    )

    doc.add_paragraph()

    # Sous-section : Interfaces publiques
    t_pub = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t_pub)
    cp = t_pub.cell(0, 0)
    _set_cell_bg(cp, '2E86C1')
    pp = cp.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp = pp.add_run('  A.  INTERFACES PUBLIQUES (sans connexion)')
    sf(rp, 12, bold=True, color=C_WHITE)
    doc.add_paragraph()

    # Capture 1 — Accueil clair
    h3('Capture 1 — Page d\'accueil (mode clair)')
    body(
        'La page d\'accueil de TuteurIA est la vitrine de l\'application. Elle presente '
        'la proposition de valeur en hero section avec un titre animé, les 12 matières '
        'disponibles, les fonctionnalités clés, le processus d\'utilisation en 4 étapes, '
        'des témoignages et un appel à l\'action. Les boutons "Commencer gratuitement" '
        'et "Se connecter" redirigent respectivement vers /auth/signup et /auth/login. '
        'La barre de navigation inclut les boutons de changement de langue (FR/EN) '
        'et de mode sombre.'
    )
    insert_screenshot('01_accueil.png', 'Page d\'accueil — Mode clair (desktop 1440px)', 12)

    # Capture 2 — Accueil sombre
    h3('Capture 2 — Page d\'accueil (mode sombre)')
    body(
        'Le mode sombre de TuteurIA utilise une palette de gris profonds (#1A1A2E, '
        '#gray-950) avec des accents bleu ciel et violet. Le basculement est instantané '
        'et persiste entre les sessions grâce au localStorage. Le mode sombre est '
        'accessible depuis l\'icône Lune/Soleil dans la Navbar sur toutes les pages.'
    )
    insert_screenshot('02_accueil_dark.png', 'Page d\'accueil — Mode sombre (desktop 1440px)', 13)

    page_break()

    # Captures auth
    t_auth = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t_auth)
    ca = t_auth.cell(0, 0)
    _set_cell_bg(ca, '1E3A5F')
    pa = ca.paragraphs[0]
    pa.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ra = pa.add_run('  B.  AUTHENTIFICATION')
    sf(ra, 12, bold=True, color=C_WHITE)
    doc.add_paragraph()

    h3('Capture 3 — Page d\'inscription')
    body(
        'La page d\'inscription (/auth/signup) présente un formulaire élégant sur fond '
        'dégradé sombre avec des blobs lumineux animés en arrière-plan. Le formulaire '
        'collecte le nom complet, l\'email et le mot de passe (min. 6 caractères, '
        'avec confirmation). Des contrôles flottants en haut à droite permettent de '
        'changer la langue et le thème sans quitter la page. La validation côté client '
        'affiche des messages d\'erreur contextuels avant l\'envoi vers Supabase Auth.'
    )
    insert_screenshot('03_inscription.png', 'Page d\'inscription — /auth/signup', 14)

    h3('Capture 4 — Page de connexion')
    body(
        'La page de connexion (/auth/login) partage le même design que l\'inscription. '
        'Elle propose un bouton "Mode démo" permettant d\'explorer l\'application sans '
        'créer de compte. En cas d\'erreur d\'authentification, un message rouge animé '
        'apparaît sous le formulaire. La session Supabase est automatiquement restaurée '
        'au rechargement de la page grâce au mécanisme onAuthStateChange.'
    )
    insert_screenshot('04_connexion.png', 'Page de connexion — /auth/login', 15)

    page_break()

    # Captures app
    t_app = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t_app)
    cd = t_app.cell(0, 0)
    _set_cell_bg(cd, '1E8449')
    pd = cd.paragraphs[0]
    pd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rd = pd.add_run('  C.  APPLICATION (espace eleve connecte)')
    sf(rd, 12, bold=True, color=C_WHITE)
    doc.add_paragraph()

    h3('Capture 5 — Tableau de bord')
    body(
        'Le tableau de bord (/dashboard) est la page d\'accueil de l\'espace connecté. '
        'Il affiche quatre cartes de statistiques globales (12 matières, 200+ QCM, '
        '50+ leçons, Tuteur 24/7), une bannière citation du jour animée, le catalogue '
        'des matières en grille, des actions rapides pour les cinq fonctionnalités '
        'principales, les QCM suggérés et une prévisualisation de la progression '
        'par matière avec barres de progression colorées.'
    )
    insert_screenshot('05_dashboard.png', 'Tableau de bord — /dashboard', 16)

    h3('Capture 6 — Catalogue des matières')
    body(
        'La page des matières (/matieres) présente les 12 disciplines du programme '
        'camerounais sous forme de cartes colorées (gradients différents par matière) '
        'avec icône, nom, nombre de chapitres et de QCM disponibles. Une barre de '
        'recherche en temps réel filtre les matières par nom. Un système de filtres '
        'permet de trier par niveau de difficulté. Le hover sur chaque carte déclenche '
        'une animation de survol fluide.'
    )
    insert_screenshot('06_matieres.png', 'Catalogue des matières — /matieres', 17)

    page_break()

    h3('Capture 7 — Détail d\'une matière (Mathématiques)')
    body(
        'La page de détail (/matieres/:id) affiche une bannière hero avec le gradient '
        'de couleur de la matière, l\'icône et les statistiques (chapitres, leçons, QCM). '
        'Deux onglets permettent de basculer entre la vue "Cours" (accordéon de chapitres '
        'avec leurs leçons) et la vue "QCM" (liste des questionnaires disponibles). '
        'Un bouton flèche gauche en haut permet de revenir au catalogue. Le composant '
        'CourseReader affiche le contenu des leçons au format textbook.'
    )
    insert_screenshot('07_cours.png', 'Détail matière Mathématiques — /matieres/maths', 18)

    h3('Capture 8 — Liste des QCM')
    body(
        'La page QCM (/qcm) présente tous les questionnaires disponibles avec un panneau '
        'de filtres (matière, difficulté, durée). Chaque carte QCM affiche le titre, '
        'la matière avec son icône colorée, le nombre de questions, la durée estimée '
        'et un badge de difficulté. Un bouton "Commencer" lance le questionnaire. '
        'Les QCM récemment complétés affichent leur score précédent.'
    )
    insert_screenshot('08_qcm_liste.png', 'Liste des QCM — /qcm', 19)

    page_break()

    # Captures IA
    t_ia = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t_ia)
    ci = t_ia.cell(0, 0)
    _set_cell_bg(ci, 'D35400')
    pi = ci.paragraphs[0]
    pi.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ri = pi.add_run('  D.  FONCTIONNALITES IA')
    sf(ri, 12, bold=True, color=C_WHITE)
    doc.add_paragraph()

    h3('Capture 9 — Tuteur IA conversationnel')
    body(
        'La page Tuteur IA (/ai-tuteur) présente une interface de chat moderne avec '
        'bulles de messages différenciées (bleue pour l\'élève, verte pour l\'IA), '
        'indicateur de frappe animé (trois points rebondissants) pendant la génération '
        'de la réponse, suggestions de questions pré-définies pour démarrer, bouton '
        'de réinitialisation de conversation et un badge "Mode démo" si la clé API '
        'Groq n\'est pas configurée. Le tuteur maintient le contexte de la conversation '
        'pour des réponses cohérentes et progressives.'
    )
    insert_screenshot('09_tuteur_ia.png', 'Tuteur IA conversationnel — /ai-tuteur', 20)

    h3('Capture 10 — Générateur Document → QCM')
    body(
        'La page Doc → QCM (/doc-quiz) propose une zone de dépôt de fichiers avec '
        'glisser-déposer (drag & drop) et sélection manuelle. Les formats supportés '
        'sont PDF, TXT, DOC et DOCX (max 10 Mo). L\'élève choisit le nombre de '
        'questions (5 à 20) et le niveau de difficulté (facile, moyen, difficile). '
        'Après analyse par l\'IA, un QCM interactif est généré avec corrections '
        'détaillées pour chaque question. Un badge orange "Clé Groq non configurée" '
        'indique le mode démo si nécessaire.'
    )
    insert_screenshot('10_doc_quiz.png', 'Générateur Doc → QCM — /doc-quiz', 21)

    page_break()

    # Captures progression + profil
    t_prog = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t_prog)
    cp2 = t_prog.cell(0, 0)
    _set_cell_bg(cp2, '6C3483')
    pp2 = cp2.paragraphs[0]
    pp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp2 = pp2.add_run('  E.  SUIVI ET PROFIL')
    sf(rp2, 12, bold=True, color=C_WHITE)
    doc.add_paragraph()

    h3('Capture 11 — Tableau de progression')
    body(
        'La page Progression (/progression) affiche trois types de graphiques interactifs '
        '(Recharts) : un graphique radar montrant le niveau par matière, un graphique '
        'linéaire de l\'évolution du score sur 8 semaines, et un graphique en barres '
        'des scores par matière. Les statistiques globales (QCM tentés, score moyen, '
        'matières travaillées, taux de réussite) sont affichées sous forme de badges. '
        'Un sélecteur d\'onglets permet de basculer entre les trois vues graphiques.'
    )
    insert_screenshot('11_progression.png', 'Tableau de progression — /progression', 22)

    h3('Capture 12 — Page profil utilisateur')
    body(
        'La page Profil (/profil) affiche les informations de l\'utilisateur connecté '
        '(nom, email, badge "Plan gratuit"), ses statistiques personnelles (QCM réalisés, '
        'score moyen, jours actifs), et ses préférences : toggle mode sombre avec '
        'animation, toggles de notifications (rappels d\'étude, nouveaux QCM). '
        'Un bouton de sauvegarde des préférences et un bouton de déconnexion '
        '(rouge avec confirmation) complètent la page.'
    )
    insert_screenshot('12_profil.png', 'Page profil utilisateur — /profil', 23)

    page_break()

    # Captures mobile
    t_mob = doc.add_table(1, 1, Cm(15.5))
    _remove_table_borders(t_mob)
    cm = t_mob.cell(0, 0)
    _set_cell_bg(cm, '1A5276')
    pm = cm.paragraphs[0]
    pm.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rm = pm.add_run('  F.  VERSION MOBILE (390 x 844px — iPhone 14)')
    sf(rm, 12, bold=True, color=C_WHITE)
    doc.add_paragraph()

    body(
        'TuteurIA est entièrement responsive. Sur mobile, la barre de navigation '
        'supérieure est remplacée par une barre d\'onglets fixe en bas de l\'écran '
        '(bottom tab bar) avec les 6 icônes de navigation. Le contenu s\'adapte '
        'automatiquement à la largeur réduite. Les cartes passent de 3 colonnes '
        '(desktop) à 2 colonnes (mobile). Les marges et espacements sont ajustés '
        'via les breakpoints Tailwind CSS (sm:, lg:).'
    )

    # Afficher les deux captures mobile cote a cote
    t_mob_shots = doc.add_table(1, 2, Cm(15.5))
    _remove_table_borders(t_mob_shots)
    t_mob_shots.alignment = WD_TABLE_ALIGNMENT.CENTER
    cm1 = t_mob_shots.cell(0, 0)
    cm2 = t_mob_shots.cell(0, 1)
    cm1.width = Cm(7.5)
    cm2.width = Cm(7.5)

    path13 = os.path.join(SHOTS, '13_mobile_accueil.png')
    path14 = os.path.join(SHOTS, '14_mobile_tuteur.png')

    if os.path.exists(path13):
        p13 = cm1.paragraphs[0]
        p13.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run13 = p13.add_run()
        run13.add_picture(path13, width=Cm(7))
    else:
        cm1.paragraphs[0].add_run('[mobile accueil]')

    if os.path.exists(path14):
        p14 = cm2.paragraphs[0]
        p14.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run14 = p14.add_run()
        run14.add_picture(path14, width=Cm(7))
    else:
        cm2.paragraphs[0].add_run('[mobile tuteur]')

    doc.add_paragraph()

    # Legendes
    t_leg = doc.add_table(1, 2, Cm(15.5))
    _remove_table_borders(t_leg)
    t_leg.alignment = WD_TABLE_ALIGNMENT.CENTER
    cl1 = t_leg.cell(0, 0)
    cl2 = t_leg.cell(0, 1)
    cl1.width = Cm(7.5)
    cl2.width = Cm(7.5)
    pl1 = cl1.paragraphs[0]
    pl1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl1 = pl1.add_run('Figure 24 : Accueil mobile (390px)')
    sf(rl1, 10, italic=True, color=C_DGRAY)
    pl2 = cl2.paragraphs[0]
    pl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl2 = pl2.add_run('Figure 25 : Tuteur IA mobile (390px)')
    sf(rl2, 10, italic=True, color=C_DGRAY)

    doc.add_paragraph()

    # ── Installation ───────────────────────────────────────────────────────────
    h2('VII.2 — Installation et configuration')
    add_table(
        ['Prerequis','Version','Verification'],
        [
            ['Node.js','18.0.0 LTS','node --version'],
            ['npm','9.0.0','npm --version'],
            ['Git','2.40+','git --version'],
            ['Navigateur','Chrome 110+','about:version'],
        ]
    )
    h3('Etapes d\'installation')
    add_table(
        ['Etape','Commande','Description'],
        [
            ['1. Cloner','git clone <url> tuteuria','Telecharger le code source'],
            ['2. Entrer','cd tuteuria','Se placer dans le dossier'],
            ['3. Installer','npm install','Installer les dependances'],
            ['4. Configurer','cp .env.example .env','Creer le fichier d\'environnement'],
            ['5. Lancer','npm run dev','Demarrer en developpement'],
            ['6. Ouvrir','http://localhost:5173','Acceder a l\'application'],
        ]
    )
    h3('Variables d\'environnement (.env)')
    add_table(
        ['Variable','Description','Source'],
        [
            ['VITE_SUPABASE_URL','URL du projet Supabase','Dashboard Supabase > Settings > API'],
            ['VITE_SUPABASE_ANON_KEY','Cle publique Supabase','Dashboard Supabase > Settings > API'],
            ['VITE_GROQ_API_KEY','Cle API Groq (tuteur IA)','console.groq.com/keys (gratuit)'],
        ], [5, 5, 4.5]
    )

    h2('VII.3 — FAQ')
    for q, r in [
        ('L\'application fonctionne-t-elle sans connexion ?',
         'Non, une connexion Internet est requise pour l\'authentification Supabase et le tuteur IA Groq. Les cours statiques sont en cache navigateur.'),
        ('Mes donnees sont-elles securisees ?',
         'Oui. JWT chiffres, RLS PostgreSQL, HTTPS obligatoire, cles API dans .env non committe.'),
        ('Le tuteur IA peut-il se tromper ?',
         'Comme tout LLM, LLaMA 3.3 peut occasionnellement produire des inexactitudes. Croiser avec les manuels officiels pour les formules importantes.'),
        ('Comment changer la langue ?',
         'Cliquer sur FR/EN dans la Navbar. Le choix est sauvegarde en localStorage.'),
        ('L\'application est-elle gratuite ?',
         'Oui, entierement gratuite. Les tiers Supabase, Groq et Vercel couvrent les couts avec leurs plans gratuits.'),
    ]:
        h4(f'Q : {q}')
        body(f'R : {r}')

    h2('VII.4 — Conclusion')
    body('TuteurIA offre une experience utilisateur moderne et intuitive, accessible sans formation prealable. L\'interface bilingue, le mode sombre et la compatibilite mobile garantissent une utilisation optimale depuis tout appareil. Le mode demo permet d\'explorer toutes les fonctionnalites sans configuration.')
    page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CONCLUSION + PERSPECTIVES
# ═══════════════════════════════════════════════════════════════════════════════

def make_conclusion():
    h1('CONCLUSION GENERALE')
    body('Le projet TuteurIA constitue une reponse concrete et ambitieuse a la problematique de l\'acces aux ressources pedagogiques numeriques pour les eleves camerounais preparant le Baccalaureat et le GCE A-Level. En mobilisant React 18, Vite 5, Tailwind CSS, Supabase et l\'API Groq, nous avons realise une plateforme complete, fonctionnelle et deployee gratuitement.')
    body('Les objectifs initiaux ont ete pleinement atteints : catalogue de 12 matieres, QCM interactifs, tuteur IA LLaMA 3.3 70B, generateur de QCM documentaire, tableau de progression, interface bilingue et mode sombre. Ce projet demontre qu\'un etudiant en Genie Logiciel peut concevoir seul une application de niveau professionnel avec les technologies open-source actuelles.')
    body('D\'un point de vue academique, ce projet a permis de consolider des competences en : architecture SPA React, Context API, integration REST et IA, modelisation UML selon 2TUP, et deploiement cloud serverless. Il illustre la capacite des technologies freemium a democratiser les solutions numeriques educatives en Afrique.')

    h1('PERSPECTIVES')
    h3('Court terme (v1.1)')
    for p in [
        'Sauvegarde des resultats QCM dans Supabase et graphiques de progression reels',
        'Protection des routes via ProtectedRoute pour /dashboard et pages authentifiees',
        'Systeme de notifications push pour les rappels de revision',
        '5 matieres supplementaires : Economie, Sciences Sociales, Chimie, Philosophie',
    ]:
        bullet(p)
    h3('Moyen terme (v2.0)')
    for p in [
        'Application mobile React Native ou PWA avec support offline',
        'Gamification : badges, classements, streaks de revision quotidienne',
        'Profils adaptatifs : le tuteur adapte son niveau selon l\'eleve',
        'Tableau de bord enseignant pour le suivi des eleves',
    ]:
        bullet(p)
    h3('Long terme (v3.0)')
    for p in [
        'Extension aux autres pays francophones (Senegal, Cote d\'Ivoire, Mali)',
        'Partenariat avec le Ministere de l\'Education du Cameroun',
        'Modele IA fine-tune sur les programmes officiels camerounais',
    ]:
        bullet(p)
    page_break()

def make_biblio():
    h1('BIBLIOGRAPHIE')
    refs = [
        '[1] JACOBSON I., BOOCH G., RUMBAUGH J. (1999). The Unified Software Development Process. Addison-Wesley.',
        '[2] ROQUES P. (2006). UML 2 par la pratique, 5e ed. Eyrolles, Paris.',
        '[3] FAVIER M. et al. (2004). Methode 2TUP. Hermes Lavoisier.',
        '[4] FREEMAN E. (2014). Head First Design Patterns. O\'Reilly Media.',
        '[5] BLOOM B.S. (1984). The 2 Sigma Problem. Educational Researcher, 13(6), 4-16.',
        '[6] UNESCO (2023). Rapport mondial de suivi sur l\'education — Technologie. Paris.',
        '[7] TOUVRON H. et al. (2023). Llama 2: Open Foundation and Fine-Tuned Chat Models. Meta AI.',
    ]
    for r in refs:
        p = doc.add_paragraph()
        run = p.add_run(r)
        sf(run, 11)
        set_spacing(p)
        p.paragraph_format.left_indent    = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)

    h1('WEBOGRAPHIE')
    urls = [
        '[W1] React 18 — https://react.dev/',
        '[W2] Vite 5 — https://vitejs.dev/',
        '[W3] Tailwind CSS — https://tailwindcss.com/docs',
        '[W4] Supabase — https://supabase.com/docs',
        '[W5] Groq API — https://console.groq.com/docs/openai',
        '[W6] Framer Motion — https://www.framer.com/motion/',
        '[W7] React Router v6 — https://reactrouter.com/',
        '[W8] Lucide React — https://lucide.dev/',
        '[W9] Recharts — https://recharts.org/',
        '[W10] Betakuma — https://www.betakuma.com',
    ]
    for u in urls:
        p = doc.add_paragraph()
        run = p.add_run(u)
        sf(run, 11)
        set_spacing(p)
        p.paragraph_format.left_indent    = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
    page_break()

def make_toc():
    h1('TABLE DES MATIERES')
    entries = [
        ('Dedicace', '3', False),
        ('Remerciements', '4', False),
        ('Resume / Abstract', '5', False),
        ('Sigles et abreviations', '7', False),
        ('Glossaire', '8', False),
        ('Introduction generale', '9', False),
        ('', '', False),
        ('DOSSIER 1 — Etude de l\'existant', '11', True),
        ('  I.1  Presentation du domaine', '11', False),
        ('  I.2  Etude de l\'existant', '13', False),
        ('  I.3  Analyse comparative', '14', False),
        ('  I.4  Limites et problematique', '15', False),
        ('  I.5  Solution proposee', '16', False),
        ('', '', False),
        ('DOSSIER 2 — Cahier des charges', '17', True),
        ('  II.1  Contexte et objectifs', '17', False),
        ('  II.2  Besoins fonctionnels', '18', False),
        ('  II.3  Besoins non fonctionnels', '20', False),
        ('  II.4  Ressources et budget', '20', False),
        ('  II.5  Diagramme de Gantt', '21', False),
        ('', '', False),
        ('DOSSIER 3 — Dossier d\'analyse', '22', True),
        ('  III.1  Methodologie 2TUP', '22', False),
        ('  III.2  Diagrammes de cas d\'utilisation', '23', False),
        ('  III.3  Diagrammes de sequence', '27', False),
        ('  III.4  Diagramme d\'activites', '31', False),
        ('', '', False),
        ('DOSSIER 4 — Dossier de conception', '32', True),
        ('  IV.1  Diagramme de classes', '32', False),
        ('  IV.2  Dictionnaire des donnees', '34', False),
        ('  IV.3  Architecture logique', '36', False),
        ('  IV.4  Diagramme d\'etat-transition', '38', False),
        ('  IV.5  Architecture physique', '39', False),
        ('', '', False),
        ('DOSSIER 5 — Realisation', '40', True),
        ('  V.1  Technologies utilisees', '40', False),
        ('  V.2  Architecture logicielle', '41', False),
        ('  V.3  Securite', '43', False),
        ('', '', False),
        ('DOSSIER 6 — Tests et validation', '44', True),
        ('  VI.1  Strategie de tests', '44', False),
        ('  VI.2  Cas de tests', '45', False),
        ('  VI.3  Performance et compatibilite', '47', False),
        ('', '', False),
        ('DOSSIER 7 — Guide utilisateur', '48', True),
        ('  VII.1  Galerie de captures d\'ecran (14 captures)', '48', False),
        ('    A. Interfaces publiques (accueil clair/sombre)', '48', False),
        ('    B. Authentification (inscription, connexion)', '50', False),
        ('    C. Application connectee (dashboard, matieres, cours, QCM)', '51', False),
        ('    D. Fonctionnalites IA (Tuteur IA, Doc → QCM)', '54', False),
        ('    E. Suivi et profil', '56', False),
        ('    F. Version mobile (390px)', '58', False),
        ('  VII.2  Installation et configuration', '59', False),
        ('  VII.3  FAQ', '61', False),
        ('', '', False),
        ('Conclusion generale', '62', False),
        ('Perspectives', '63', False),
        ('Bibliographie', '64', False),
        ('Webographie', '65', False),
    ]
    for entry, page, bold in entries:
        if not entry:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p.add_run(entry)
        sf(run1, 12 if bold else 11, bold=bold, color=C_BLUE if bold else C_BLACK)
        run2 = p.add_run(f'{"." * max(1, 70 - len(entry) - len(page))} {page}')
        sf(run2, 10, color=C_DGRAY)
        set_spacing(p, after=2)

# ─── BUILD ────────────────────────────────────────────────────────────────────
print('Generation du rapport TuteurIA v2 (avec header/footer/pagination)...')
make_cover()
make_dedicace()
make_remerciements()
make_resume()
make_sigles()
make_glossaire()
make_intro()
make_dossier1()
make_dossier2()
make_dossier3()
make_dossier4()
make_dossier5()
make_dossier6()
make_dossier7()
make_conclusion()
make_biblio()
make_toc()

doc.save(OUT)
print(f'Rapport sauvegarde : {OUT}')

from docx import Document as D2
d2 = D2(OUT)
print(f'Paragraphes : {len(d2.paragraphs)}')
print(f'Tableaux    : {len(d2.tables)}')
nimg = sum(1 for p in d2.paragraphs for r in p.runs
           if r._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'))
print(f'Images      : {nimg}')
